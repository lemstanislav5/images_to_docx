import os
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image


def create_documents(image_folder, output_docx, index_docx, image_width=6.0):
    # Создаем основной документ
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = 14

    # Создаем документ для оглавления
    index_doc = Document()
    style = index_doc.styles['Normal']
    index_doc.add_paragraph('Оглавление изображений')
    
    # Добавляем таблицу в оглавление (2 колонки: имя файла и страница)
    table = index_doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Имя файла'
    hdr_cells[1].text = 'Страница фототаблицы'
    
    # Получаем список изображений в папке
    image_files = [f for f in os.listdir(image_folder) 
                  if f.lower().endswith(('.png', '.jpg', '.jpeg', '.gif', '.bmp'))]
    print(f"количество изображений: {len(image_files)}")

    # Добавляем оглавление
    title = doc.add_paragraph('Фототаблица')  
    # Уровень 0 — самый высокий, обычно используется для заголовков документов  
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Выравнивание по центру  
    for page_num, image_file in enumerate(image_files, start=1):
        image_path = os.path.join(image_folder, image_file)
        
        try:
            # Добавляем номер страницы под изображением
            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if page_num != 1: 
                run = paragraph.add_run(f'{page_num}')
            else: 
                run = paragraph.add_run(f'')
            run.bold = False

            # Добавляем изображение в основной документ
            doc.add_picture(image_path, width=Inches(image_width))
                        
            # Добавляем разрыв страницы (кроме последней)
            if page_num < len(image_files):
                doc.add_page_break()
            
            # Добавляем запись в оглавление
            row_cells = table.add_row().cells
            row_cells[0].text = image_file
            row_cells[1].text = str(page_num)
            
        except Exception as e:
            print(f"Ошибка при обработке файла {image_file}: {e}")
    
    # Сохраняем документы
    doc.save(output_docx)
    index_doc.save(index_docx)
    print(f"Созданы документы: {output_docx} и {index_docx}")

# Пример использования
if __name__ == "__main__":
    image_folder = 'images'  # Папка с изображениями
    output_docx = 'Фототаблица.docx'  # Основной документ
    index_docx = 'Список изображений.docx'     # Документ с оглавлением
    
    # Создаем папку, если её нет
    if not os.path.exists(image_folder):
        os.makedirs(image_folder)
        print(f"Создана папка '{image_folder}'. Добавьте в неё изображения и запустите скрипт снова.")
    else:
        create_documents(image_folder, output_docx, index_docx, image_width=6.0)
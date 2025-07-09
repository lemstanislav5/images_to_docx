import os
import tempfile
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from PIL import Image, UnidentifiedImageError
import sys
import hashlib
from tqdm import tqdm
import tkinter as tk
from tkinter import filedialog, ttk, messagebox

class PhotoTableApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Создание фототаблицы")
        self.root.geometry("650x300")  # Увеличена ширина окна на 1 см (было 600)
        
        # Переменные для хранения путей
        self.source_folder = tk.StringVar()
        self.output_folder = tk.StringVar()
        
        # Создаем элементы интерфейса
        self.create_widgets()
        
    def create_widgets(self):
        # Фрейм для основной части
        main_frame = ttk.Frame(self.root, padding="10")
        main_frame.pack(fill=tk.BOTH, expand=True)
        
        # Выбор папки с изображениями
        ttk.Label(main_frame, text="Папка с изображениями:").grid(row=0, column=0, sticky=tk.W, pady=5)
        source_entry = ttk.Entry(main_frame, textvariable=self.source_folder, width=50)
        source_entry.grid(row=0, column=1, padx=5)
        ttk.Button(main_frame, text="Обзор...", command=self.browse_source).grid(row=0, column=2, padx=5)
        
        # Выбор папки для сохранения
        ttk.Label(main_frame, text="Папка для сохранения документов:").grid(row=1, column=0, sticky=tk.W, pady=5)
        output_entry = ttk.Entry(main_frame, textvariable=self.output_folder, width=50)
        output_entry.grid(row=1, column=1, padx=5)
        ttk.Button(main_frame, text="Обзор...", command=self.browse_output).grid(row=1, column=2, padx=5)
        
        # Кнопка старта - выровнена по левому краю (column=0)
        ttk.Button(main_frame, text="Создать документы", command=self.start_processing).grid(
            row=2, column=0, pady=20, sticky=tk.W)
        
        # Прогресс-бар
        self.progress = ttk.Progressbar(main_frame, orient=tk.HORIZONTAL, length=400, mode='determinate')
        self.progress.grid(row=3, column=0, columnspan=3, pady=10)
        
        # Статус
        self.status_label = ttk.Label(main_frame, text="Готов к работе")
        self.status_label.grid(row=6, column=0, columnspan=3)
    
    def browse_source(self):
        folder = filedialog.askdirectory(title="Выберите папку с изображениями")
        if folder:
            self.source_folder.set(folder)
    
    def browse_output(self):
        folder = filedialog.askdirectory(title="Выберите папку для сохранения документов")
        if folder:
            self.output_folder.set(folder)
    
    def update_progress(self, value):
        self.progress['value'] = value
        self.root.update_idletasks()
    
    def show_message(self, title, message):
        messagebox.showinfo(title, message)
    
    def get_unique_image_cache_name(self, image_path):
        """Генерирует уникальное имя для кэша изображений"""
        with open(image_path, 'rb') as f:
            file_hash = hashlib.md5(f.read()).hexdigest()
        return f"img_{file_hash}"
    
    def is_image_valid(self, image_path):
        """Проверяет, можно ли открыть изображение"""
        try:
            with Image.open(image_path) as img:
                img.verify()
            return True
        except (IOError, SyntaxError, UnidentifiedImageError) as e:
            print(f"\nПоврежденное изображение {os.path.basename(image_path)}: {str(e)}")
            return False
    
    def calculate_image_size(self, img_path, max_width=Inches(6.0), max_height=Inches(8.0)):
        """Рассчитывает размер изображения с сохранением пропорций"""
        try:
            with Image.open(img_path) as img:
                width, height = img.size
                ratio = min(max_width / width, max_height / height)
                return width * ratio, height * ratio
        except Exception as e:
            print(f"\nОшибка при расчете размера {os.path.basename(img_path)}: {str(e)}")
            return max_width, max_height
    
    def add_header_with_page_num(self, section):
        """Добавляет номер страницы в верхний колонтитул, начиная со 2 страницы"""
        header = section.header
        for elem in header._element:
            header._element.remove(elem)
        
        paragraph = header.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        fld_char1 = OxmlElement('w:fldChar')
        fld_char1.set(qn('w:fldCharType'), 'begin')
        
        instr_text = OxmlElement('w:instrText')
        instr_text.text = 'IF  \\* MERGEFORMAT 1 < > 1  "Страница { PAGE }" ""'
        
        fld_char2 = OxmlElement('w:fldChar')
        fld_char2.set(qn('w:fldCharType'), 'end')
        
        run = paragraph.add_run()
        run._r.append(fld_char1)
        run._r.append(instr_text)
        run._r.append(fld_char2)
    
    def start_processing(self):
        source_folder = self.source_folder.get()
        output_folder = self.output_folder.get()
        
        if not source_folder or not output_folder:
            self.show_message("Ошибка", "Пожалуйста, укажите обе папки")
            return
        
        if not os.path.exists(source_folder):
            self.show_message("Ошибка", f"Папка '{source_folder}' не найдена!")
            return
        
        try:
            image_files = sorted([
                f for f in os.listdir(source_folder) 
                if f.lower().endswith(('.png', '.jpg', '.jpeg', '.bmp', '.heic'))
            ])
        except Exception as e:
            self.show_message("Ошибка", f"Ошибка при чтении папки: {str(e)}")
            return
        
        if not image_files:
            self.show_message("Ошибка", "В указанной папке нет изображений!")
            return
        
        self.status_label.config(text=" Обработка...  ")
        self.progress['maximum'] = len(image_files)
        self.progress['value'] = 0
        self.root.update_idletasks()
        
        doc = Document()
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(12)
        
        # Устанавливаем поля страницы
        for section in doc.sections:
            section.top_margin = Cm(1.5)
            section.bottom_margin = Cm(1.5)
            section.left_margin = Cm(1.5)
            section.right_margin = Cm(1.5)
        
        self.add_header_with_page_num(doc.sections[0])
        
        # Документ для оглавления
        index_doc = Document()
        style_index_doc = index_doc.styles['Normal']
        font_index_doc = style_index_doc.font
        font_index_doc.name = 'Arial'
        font_index_doc.size = Pt(12)
        index_doc.add_paragraph('Оглавление изображений', style='Heading1')
        table = index_doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Имя файла'
        hdr_cells[1].text = 'Страница'
        
        # Титульная страница
        title = doc.add_paragraph('Фототаблица')
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        successful_images = 0
        total_images = len(image_files)
        
        for page_num, image_file in enumerate(image_files, start=1):
            image_path = os.path.join(source_folder, image_file)
            
            if not self.is_image_valid(image_path):
                self.update_progress(page_num)
                continue
                
            try:
                # Создаем временную копию изображения с уникальным именем
                with tempfile.NamedTemporaryFile(suffix=os.path.splitext(image_file)[1], delete=False) as tmp_file:
                    with open(image_path, 'rb') as original_file:
                        tmp_file.write(original_file.read())
                    tmp_path = tmp_file.name
                
                img_width, img_height = self.calculate_image_size(tmp_path)

                # Добавляем номер страницы под изображением
                paragraph = doc.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                if page_num != 1: 
                    run = paragraph.add_run(f'{page_num}')
                else: 
                    run = paragraph.add_run(f'')
                run.bold = False

                # Добавляем изображение через временный файл
                paragraph = doc.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = paragraph.add_run()
                
                try:
                    run.add_picture(tmp_path, width=img_width, height=img_height)
                    os.unlink(tmp_path)  # Удаляем временный файл
                    successful_images += 1  # Переносим увеличение счетчика сюда
                except Exception as e:
                    print(f"\nОшибка при добавлении {image_file}: {str(e)}")
                    if os.path.exists(tmp_path):
                        os.unlink(tmp_path)
                    self.update_progress(page_num)
                    continue
                
                # Добавляем запись в оглавление
                row_cells = table.add_row().cells
                row_cells[0].text = image_file
                row_cells[1].text = str(page_num)
                
                if page_num < len(image_files) + 1:
                    doc.add_page_break()
                    
            except Exception as e:
                print(f"\nОшибка обработки {image_file}: {str(e)}")
                if 'tmp_path' in locals() and os.path.exists(tmp_path):
                    os.unlink(tmp_path)
                self.update_progress(page_num)
                continue
            
            self.update_progress(page_num)
        
        try:
            output_docx = os.path.join(output_folder, 'Фототаблица.docx')
            index_docx = os.path.join(output_folder, 'Оглавление.docx')
            
            doc.save(output_docx)
            index_doc.save(index_docx)
            
            # Исправленный вывод информации о количестве файлов
            self.status_label.config(text=f"Готово! Обработано: {successful_images}/{total_images}")
            self.show_message("Готово", 
                f"Всего изображений: {total_images}\n"
                f"Успешно добавлено: {successful_images}\n"
                f"Созданы:\n- {output_docx}\n- {index_docx}")
        except Exception as e:
            self.status_label.config(text="Ошибка сохранения")
            self.show_message("Ошибка", f"Ошибка сохранения: {str(e)}")

if __name__ == "__main__":
    root = tk.Tk()
    app = PhotoTableApp(root)
    root.mainloop()